#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def add_hyperlink(paragraph, text, url):
    """段落にハイパーリンクを追加する"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # ハイパーリンクのスタイル設定
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    run.append(rPr)
    run.text = text
    hyperlink.append(run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def setup_document_styles(doc):
    """文書のスタイルを設定する"""
    # 標準スタイルの設定
    styles = doc.styles
    
    # 見出し1のスタイル
    heading1 = styles['Heading 1']
    heading1.font.name = 'Yu Gothic'
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    
    # 見出し2のスタイル
    heading2 = styles['Heading 2']
    heading2.font.name = 'Yu Gothic'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    
    # 見出し3のスタイル
    heading3 = styles['Heading 3']
    heading3.font.name = 'Yu Gothic'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    
    # 標準テキストのスタイル
    normal = styles['Normal']
    normal.font.name = 'Yu Gothic'
    normal.font.size = Pt(10)
    
    return doc

def parse_markdown_table(lines, start_idx):
    """マークダウンテーブルを解析する"""
    table_lines = []
    i = start_idx
    
    while i < len(lines):
        line = lines[i].strip()
        if not line or not line.startswith('|'):
            break
        if not line.startswith('|---') and not line.startswith('|-'):
            table_lines.append(line)
        i += 1
    
    if not table_lines:
        return None, start_idx
    
    # テーブルデータを解析
    table_data = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            table_data.append(cells)
    
    return table_data, i

def markdown_to_docx(md_file, docx_file):
    """マークダウンファイルをWord文書に変換する"""
    
    # マークダウンファイルを読み込み
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 新しいWord文書を作成
    doc = Document()
    doc = setup_document_styles(doc)
    
    # ページ設定
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)    # A4
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # 見出し処理
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # リスト処理
        elif line.startswith('- '):
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            # 太字とマークダウン記法を処理
            text = line[2:]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 太字マークダウンを除去
            run = p.add_run(text)
            run.font.name = 'Yu Gothic'
            run.font.size = Pt(10)
            
            # 太字部分を再度処理
            if '**' in line:
                p.clear()
                parts = re.split(r'(\*\*.*?\*\*)', line[2:])
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.name = 'Yu Gothic'
                    run.font.size = Pt(10)
        
        # テーブル処理
        elif line.startswith('|'):
            table_data, next_i = parse_markdown_table(lines, i)
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = cell_data
                        
                        # ヘッダー行の場合は太字
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.name = 'Yu Gothic'
                                    run.font.size = Pt(9)
                        else:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Yu Gothic'
                                    run.font.size = Pt(9)
                
                i = next_i - 1
        
        # コードブロック処理
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                # 背景色を設定（薄いグレー）
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.space_before = Pt(6)
        
        # 通常の段落
        else:
            if line:
                p = doc.add_paragraph()
                
                # 太字とマークダウン記法を処理
                if '**' in line:
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            run = p.add_run(part)
                        run.font.name = 'Yu Gothic'
                        run.font.size = Pt(10)
                else:
                    run = p.add_run(line)
                    run.font.name = 'Yu Gothic'
                    run.font.size = Pt(10)
        
        i += 1
    
    # 文書を保存
    doc.save(docx_file)
    print(f"✓ Word文書が作成されました: {docx_file}")

def main():
    """メイン関数"""
    md_files = [
        "HM_スキルシート.md",
        "README.md"
    ]
    
    print("マークダウンからWord文書への変換を開始します...\n")
    
    for md_file in md_files:
        if os.path.exists(md_file):
            docx_file = md_file.replace('.md', '.docx')
            try:
                markdown_to_docx(md_file, docx_file)
                print(f"✓ {md_file} → {docx_file}")
            except Exception as e:
                print(f"✗ エラー: {md_file} の変換に失敗しました - {e}")
        else:
            print(f"✗ ファイルが見つかりません: {md_file}")
    
    print("\n変換処理が完了しました！")
    print("\n作成されたWord文書の特徴:")
    print("- 日本語フォント（Yu Gothic）を使用")
    print("- 適切な見出しレベルとスタイル")
    print("- 箇条書きとテーブルの自動変換")
    print("- A4サイズの印刷用レイアウト")

if __name__ == "__main__":
    main() 